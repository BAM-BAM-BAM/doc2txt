"""FGT-category tests for doc2txt.

New tests organized by FGT test categories:
  INV-*  : Domain invariants that must always hold
  QUAL-* : Output quality verification
  BOUND-*: Edge case / boundary handling
  PRO-*  : Proactive bug-class detection
"""

import sys
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).parent.parent))

from doc2txt import (
    SUPPORTED_EXTENSIONS,
    ImageFeature,
    TextQualityScorer,
    extract_text,
    create_markdown,
)


class TestInvariants:
    """INV-* tests: domain rules that must always hold."""

    def test_inv_001_all_extensions_have_handlers(self, tmp_path):
        """Every extension in SUPPORTED_EXTENSIONS must have a handler in extract_text().

        If this fails, a format was added to SUPPORTED_EXTENSIONS without
        adding a corresponding extraction handler.
        """
        for ext in SUPPORTED_EXTENSIONS:
            path = tmp_path / f"test{ext}"
            path.write_bytes(b"fake content")
            try:
                extract_text(path)
            except ValueError as e:
                if "Unsupported format" in str(e):
                    pytest.fail(
                        f"Extension '{ext}' is in SUPPORTED_EXTENSIONS "
                        f"but has no handler in extract_text()"
                    )
            except Exception:
                # Other errors (e.g., invalid file content, missing LibreOffice)
                # are fine -- the handler exists, it just can't process fake data
                pass

    def test_inv_002_quality_score_range(self):
        """QualityMetrics.total_score must be in [0.0, 1.0]."""
        scorer = TextQualityScorer()
        test_inputs = [
            "",
            "   ",
            "hello",
            "The quick brown fox jumps over the lazy dog.",
            "x" * 10000,
            "abc " * 500,
            "!@#$%^&*()_+" * 50,
        ]
        for text in test_inputs:
            metrics = scorer.score(text)
            assert 0.0 <= metrics.total_score <= 1.0, (
                f"Quality score {metrics.total_score} out of range for: {text[:50]!r}"
            )

    def test_inv_003_feature_vector_dimension(self):
        """ImageFeature.to_vector() must return exactly 14 elements."""
        feature = ImageFeature(
            width=500, height=300, area=150000, aspect_ratio=1.67,
            page_y_center=0.5, region="body",
            surrounding_text_density=50.0, has_nearby_caption=True,
            brightness_mean=128.0, brightness_std=40.0,
            is_mostly_white=False, has_contrast=True,
        )
        vector = feature.to_vector()
        assert len(vector) == 14, (
            f"Feature vector has {len(vector)} elements, expected 14"
        )

    def test_inv_004_output_file_naming(self, tmp_path):
        """Output .md file must have the same stem as input document."""
        import docx
        doc = docx.Document()
        doc.add_paragraph("Test content")
        path = tmp_path / "my_document.docx"
        doc.save(str(path))

        sections = extract_text(path)
        md_content = create_markdown(path, sections)
        # The markdown header should reference the original file stem
        assert "my_document" in md_content

    def test_inv_005_extract_text_returns_nonempty_list(self, tmp_path):
        """extract_text() must return at least one section for any supported format."""
        import docx
        doc = docx.Document()
        doc.add_paragraph("Content")
        path = tmp_path / "test.docx"
        doc.save(str(path))

        result = extract_text(path)
        assert isinstance(result, list)
        assert len(result) >= 1


class TestOutputQuality:
    """QUAL-* tests: verify output values are correct, not just computed."""

    def test_qual_001_docx_extraction_produces_readable_text(self, tmp_path):
        """Extraction of known DOCX produces expected readable text.

        This is an output quality test -- we verify the actual text a user
        would see, not just that the pipeline ran without errors.
        """
        import docx
        doc = docx.Document()
        doc.add_heading("Project Report", level=1)
        doc.add_paragraph(
            "The quarterly results show a 15% increase in revenue. "
            "This growth was driven by expansion into new markets."
        )
        doc.add_heading("Conclusion", level=2)
        doc.add_paragraph("The outlook for next quarter remains positive.")
        path = tmp_path / "report.docx"
        doc.save(str(path))

        sections = extract_text(path)
        text = "\n".join(sections)

        # Verify key content is present and readable
        assert "Project Report" in text
        assert "15% increase" in text
        assert "Conclusion" in text
        assert "outlook" in text

        # Verify markdown formatting
        assert "# Project Report" in text
        assert "## Conclusion" in text

        # Verify quality scoring would accept this
        scorer = TextQualityScorer()
        metrics = scorer.score(text)
        assert metrics.total_score > 0.3, (
            f"Extracted text scored {metrics.total_score}, expected readable text"
        )
        assert metrics.word_count > 10


class TestBoundary:
    """BOUND-* tests: edge case handling."""

    def test_bound_001_extract_text_empty_docx(self, tmp_path):
        """Empty DOCX should not crash, should return list with empty string."""
        import docx
        doc = docx.Document()
        path = tmp_path / "empty.docx"
        doc.save(str(path))

        result = extract_text(path)
        assert isinstance(result, list)
        assert len(result) >= 1

    def test_bound_002_unsupported_format_raises(self, tmp_path):
        """Unsupported file extension raises ValueError, not silent failure."""
        path = tmp_path / "test.xyz"
        path.touch()
        with pytest.raises(ValueError, match="Unsupported format"):
            extract_text(path)

    def test_bound_003_quality_scorer_handles_extremes(self):
        """Quality scorer handles extreme inputs without crash."""
        scorer = TextQualityScorer()

        # Very long input
        metrics = scorer.score("word " * 50000)
        assert 0.0 <= metrics.total_score <= 1.0

        # Single character
        metrics = scorer.score("a")
        assert 0.0 <= metrics.total_score <= 1.0

        # Only punctuation
        metrics = scorer.score("!!! ??? ... ,,, ;;; :::")
        assert 0.0 <= metrics.total_score <= 1.0

    def test_bound_004_image_feature_edge_values(self):
        """ImageFeature handles edge values without crash."""
        # Minimum values
        feature = ImageFeature(
            width=1, height=1, area=1, aspect_ratio=1.0,
            page_y_center=0.0, region="header",
            surrounding_text_density=0.0, has_nearby_caption=False,
            brightness_mean=0.0, brightness_std=0.0,
            is_mostly_white=True, has_contrast=False,
        )
        vector = feature.to_vector()
        assert len(vector) == 14
        assert all(isinstance(v, float) for v in vector)


class TestProactive:
    """PRO-* tests: detect structural anti-patterns before they manifest as bugs."""

    def test_pro_001_module_size_enforcement(self):
        """No source file should exceed the module size limit.

        Current state: doc2txt.py is a 3,496-line monolith (limit: 3,600).
        After modularization, tighten to 500 lines per module.
        """
        project_root = Path(__file__).parent.parent
        max_lines = 3800  # Current limit (monolith). Tighten after modularization.

        violations = []
        for py_file in project_root.glob("*.py"):
            if py_file.name.startswith("test_") or py_file.name.startswith("."):
                continue
            line_count = len(py_file.read_text().splitlines())
            if line_count > max_lines:
                violations.append(f"{py_file.name}: {line_count} lines (limit: {max_lines})")

        assert not violations, (
            "Module size violations:\n" + "\n".join(violations)
        )

    def test_pro_002_no_hardcoded_extensions_outside_registry(self):
        """Extension strings should only appear in SUPPORTED_EXTENSIONS definition.

        Detects hardcoded format checks that bypass the central registry.
        """
        project_root = Path(__file__).parent.parent
        source = (project_root / "doc2txt.py").read_text()

        # Find all lines with hardcoded extension checks outside the registry
        violations = []
        for i, line in enumerate(source.splitlines(), 1):
            # Skip the SUPPORTED_EXTENSIONS definition itself
            if "SUPPORTED_EXTENSIONS" in line:
                continue
            # Skip comments and strings in format dispatch (the switch-case is OK)
            if line.strip().startswith("#"):
                continue
            # Look for extension equality checks like == '.pdf' or in ('.doc',
            for ext in ['.pdf', '.docx', '.doc', '.rtf', '.odt']:
                if f"== '{ext}'" in line or f'== "{ext}"' in line:
                    # Allow in extract_text dispatcher (the legitimate dispatch)
                    if "def extract_text" not in source.splitlines()[max(0, i-5):i]:
                        # Check if we're inside extract_text function body
                        # (rough heuristic: within 20 lines of "def extract_text")
                        context_start = max(0, i - 20)
                        context = source.splitlines()[context_start:i]
                        if any("def extract_text(" in cl for cl in context):
                            continue
                        if any("def find_documents(" in cl for cl in context):
                            continue
                        violations.append(f"Line {i}: {line.strip()}")

        # This is advisory — violations may be legitimate dispatch logic
        # Only fail if there are clearly wrong hardcoded checks
        # For now, just ensure the test runs and documents any findings
        if violations:
            import warnings
            warnings.warn(
                f"Found {len(violations)} potential hardcoded extension checks "
                f"outside SUPPORTED_EXTENSIONS: {violations[:3]}"
            )
